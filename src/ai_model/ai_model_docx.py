from utils.utils_file import save_file_txt, save_file_xlsx, save_file_docx
from ai_model.ai_model_pdf import PDFJPG, PDFPNG, PDFSVG
from ai_model.ai_model_translate import UniversalTranslator

import subprocess
import os

from docx import Document

from transformers import BartTokenizer, BartForConditionalGeneration

class DOCXPDF:
    def __init__(self):
        pass

    def convert(self, input_docx_path, output_pdf_path):
        try:
            # Verifica che il file di input sia un file DOCX valido
            if not input_docx_path.endswith('.docx'):
                raise Exception(f"Il file {input_docx_path} non è un file DOCX valido.")

            # Converti DOCX a PDF usando LibreOffice
            command = [
                "libreoffice",
                "--headless",  # Run without GUI
                "--convert-to", "pdf",  # Convert to PDF
                "--outdir", os.path.dirname(output_pdf_path),  # Output directory
                input_docx_path
            ]
            
            subprocess.run(command, check=True)

            # Rinominare il file PDF convertito se necessario
            generated_pdf_path = os.path.join(
                os.path.dirname(output_pdf_path), 
                os.path.splitext(os.path.basename(input_docx_path))[0] + ".pdf"
            )
            os.rename(generated_pdf_path, output_pdf_path)
            
            print(f"Conversione completata: {output_pdf_path}")
            return 0
        
        except subprocess.CalledProcessError as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1
        
        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1

class DOCXJPG:
    def __init__(self):
        # Inizializza le classi per la conversione DOCX -> PDF e PDF -> JPG
        self.docx_to_pdf_converter = DOCXPDF()
        self.pdf_to_jpg_converter = PDFJPG()

    def convert(self, input_docx_path, output_jpg_path, page_number):
        try:
            # Definisci i nomi temporanei per i file
            temp_pdf_path = "temp_file/" + input_docx_path + '.pdf'
            temp_jpg_path = output_jpg_path  # Il percorso finale del JPG

            # Converti il DOCX in PDF
            if self.docx_to_pdf_converter.convert(input_docx_path, temp_pdf_path) != 0:
                raise Exception(f"Errore durante la conversione del file DOCX in PDF: {input_docx_path}")

            # Converti il PDF in JPG
            if self.pdf_to_jpg_converter.convert(temp_pdf_path, temp_jpg_path, page_number) != 0:
                raise Exception(f"Errore durante la conversione del file PDF in JPG: {temp_pdf_path}")

            print(f"Conversione completata: {output_jpg_path}")
            return 0

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1
        
        finally:
            # Rimuove il file PDF temporaneo se esiste
            if os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                    print(f"File PDF temporaneo rimosso: {temp_pdf_path}")
                except Exception as e:
                    print(f"Errore durante la rimozione del file PDF temporaneo: {e}")

class DOCXPNG:
    def __init__(self):
        # Inizializza le classi per la conversione DOCX -> PDF e PDF -> PNG
        self.docx_to_pdf_converter = DOCXPDF()
        self.pdf_to_png_converter = PDFPNG()

    def convert(self, input_docx_path, output_png_path, page_number):
        try:
            # Definisci i nomi temporanei per i file
            temp_pdf_path = "temp_file/" + input_docx_path + '.pdf'
            temp_png_path = output_png_path  # Il percorso finale del PNG

            # Converti il DOCX in PDF
            if self.docx_to_pdf_converter.convert(input_docx_path, temp_pdf_path) != 0:
                raise Exception(f"Errore durante la conversione del file DOCX in PDF: {input_docx_path}")

            # Converti il PDF in PNG
            if self.pdf_to_png_converter.convert(temp_pdf_path, temp_png_path, page_number) != 0:
                raise Exception(f"Errore durante la conversione del file PDF in PNG: {temp_pdf_path}")

            print(f"Conversione completata: {output_png_path}")
            return 0

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1
        
        finally:
            # Rimuove il file PDF temporaneo se esiste
            if os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                    print(f"File PDF temporaneo rimosso: {temp_pdf_path}")
                except Exception as e:
                    print(f"Errore durante la rimozione del file PDF temporaneo: {e}")

class DOCXXLSX:
    def __init__(self):
        pass

    def convert(self, input_docx_path, output_xlsx_path):
        try:
            # Carica il documento DOCX
            document = Document(input_docx_path)

            # Crea una lista di righe da inserire nel file XLSX
            data = []

            # Itera sui paragrafi del documento DOCX e aggiungi ogni paragrafo come una riga della lista
            for paragraph in document.paragraphs:
                # Aggiungi il testo del paragrafo come una lista (anche se è una singola cella)
                data.append([paragraph.text])

            # Salva i dati nel file XLSX utilizzando la funzione save_file_xlsx
            if save_file_xlsx(data, output_xlsx_path) == 0:
                print(f"Conversione completata: {output_xlsx_path}")
                return 0
            else:
                raise Exception(f"Errore durante il salvataggio del file XLSX: {output_xlsx_path}")

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1

class DOCXTXT:
    def __init__(self):
        pass

    def convert(self, input_docx_path, output_txt_path):
        try:
            # Carica il file DOCX
            document = Document(input_docx_path)

            # Estrai il testo dal file DOCX
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)

            # Combina il testo in un'unica stringa separata da nuove linee
            text = "\n".join(full_text)

            # Salva il testo estratto nel file TXT
            if save_file_txt(text, output_txt_path) == 0:
                print(f"Conversione completata: {output_txt_path}")
                return 0
            else:
                raise Exception(f"Errore durante il salvataggio del file TXT: {output_txt_path}")

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1

class DOCXSVG:
    def __init__(self):
        # Inizializza le classi per la conversione DOCX -> PDF e PDF -> SVG
        self.docx_to_pdf_converter = DOCXPDF()
        self.pdf_to_svg_converter = PDFSVG()

    def convert(self, input_docx_path, output_svg_path, page_number):
        try:
            # Definisci i nomi temporanei per i file
            temp_pdf_path = "temp_file/" + input_docx_path + '.pdf'
            temp_svg_path = output_svg_path  # Il percorso finale del SVG

            # Converti il DOCX in PDF
            if self.docx_to_pdf_converter.convert(input_docx_path, temp_pdf_path) != 0:
                raise Exception(f"Errore durante la conversione del file DOCX in PDF: {input_docx_path}")

            # Converti il PDF in SVG
            if self.pdf_to_svg_converter.convert_to_svg(temp_pdf_path, temp_svg_path, page_number) != 0:
                raise Exception(f"Errore durante la conversione del file PDF in SVG: {temp_pdf_path}")

            print(f"Conversione completata: {output_svg_path}")
            return 0

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1
        
        finally:
            # Rimuove il file PDF temporaneo se esiste
            if os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                    print(f"File PDF temporaneo rimosso: {temp_pdf_path}")
                except Exception as e:
                    print(f"Errore durante la rimozione del file PDF temporaneo: {e}")

class DOCXTXTSummary:
    def __init__(self):
        # Inizializza il modello e il tokenizer per il riassunto
        self.tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
        self.model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")

    def convert(self, input_docx_path, output_txt_path):
        try:
            # Carica il file DOCX
            document = Document(input_docx_path)

            # Estrai il testo dal file DOCX
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)

            # Combina il testo in un'unica stringa separata da nuove linee
            text = "\n".join(full_text)

            # Preprocessa il testo per il modello
            inputs = self.tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)

            # Genera il riassunto
            summary_ids = self.model.generate(
                inputs["input_ids"],
                max_length=150,
                min_length=40,
                length_penalty=2.0,
                num_beams=4,
                early_stopping=True
            )

            # Decodifica il riassunto generato
            summary = self.tokenizer.decode(summary_ids[0], skip_special_tokens=True)

            # Salva il riassunto nel file TXT
            if save_file_txt(summary, output_txt_path) == 0:
                print(f"Riassunto completato e salvato in: {output_txt_path}")
                return 0
            else:
                raise Exception(f"Errore durante il salvataggio del file TXT: {output_txt_path}")

        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1
        
class DOCXTranslate:
    def __init__(self, source_lang, target_lang):
        self.translator = UniversalTranslator(source_lang, target_lang)

    def convert(self, input_docx_path, output_docx_path):
        try:
            # Carica il file DOCX
            document = Document(input_docx_path)

            # Estrai il testo dal file DOCX
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)

            # Combina il testo in un'unica stringa separata da nuove linee
            text = "\n".join(full_text)

            # Traduci il testo
            translated_text = self.translator.convert(text)

            # Salva il testo tradotto nel file DOCX
            if save_file_docx(translated_text, output_docx_path) == 0:
                print(f"Traduzione completata e salvata in: {output_docx_path}")
                return 0
            else:
                raise Exception(f"Errore durante il salvataggio del file DOCX: {output_docx_path}")
        except Exception as e:
            print(f"Errore durante la conversione del file {input_docx_path}: {e}")
            return 1