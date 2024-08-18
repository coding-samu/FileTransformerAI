from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import openpyxl
from pydub import AudioSegment

def load_file_txt(file_path):
    # verifica che l'estensione del file sia .txt
    if not file_path.endswith('.txt'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Errore durante il caricamento del file {file_path}: {e}")
        return None

def save_file_txt(data, file_path):
    # verifica che l'estensione del file sia .txt
    if not file_path.endswith('.txt'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(data)
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")
        return None
    
def load_file_pdf(file_path):
    # verifica che l'estensione del file sia .pdf
    if not file_path.endswith('.pdf'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        reader = PdfReader(file_path)
        return reader
    except Exception as e:
        print(f"Errore durante il caricamento del file {file_path}: {e}")
        return None
    
def save_file_pdf(pdf_writer, file_path):
    if not file_path.endswith('.pdf'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    if pdf_writer is None:
        print(f"Errore: il pdf_writer è None, non posso salvare il file.")
        return None
    try:
        with open(file_path, 'wb') as file:
            pdf_writer.write(file)
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")
        return None
    
def load_file_jpg(file_path):
    # verifica che l'estensione del file sia .jpg
    if not file_path.endswith('.jpg'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'rb') as file:
            return file.read()
    except Exception as e:
        print(f"Errore durante il caricamento del file {file_path}: {e}")
        return None
    
def save_file_jpg(data, file_path):
    # verifica che l'estensione del file sia .jpg
    if not file_path.endswith('.jpg'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'wb') as file:
            file.write(data)
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")
        return None
    
def load_file_png(file_path):
    # verifica che l'estensione del file sia .png
    if not file_path.endswith('.png'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'rb') as file:
            return file.read()
    except Exception as e:
        print(f"Errore durante il caricamento del file {file_path}: {e}")
        return None
    
def save_file_png(data, file_path):
    # verifica che l'estensione del file sia .png
    if not file_path.endswith('.png'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        with open(file_path, 'wb') as file:
            file.write(data)
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")
        return None
    
def load_file_docx(file_path):
    # verifica che l'estensione del file sia .docx
    if not file_path.endswith('.docx'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Errore durante la lettura del file {file_path}: {e}")
        return None
    
def save_file_docx(data, file_path):
    # verifica che l'estensione del file sia .docx
    if not file_path.endswith('.docx'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        doc = Document()
        for line in data.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
        print(f"File salvato come {file_path}")
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")

def load_file_xlsx(file_path):
    # verifica che l'estensione del file sia .xlsx
    if not file_path.endswith('.xlsx'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        data = []

        for row in sheet.iter_rows(values_only=True):
            data.append(list(row))

        return data
    except Exception as e:
        print(f"Errore durante la lettura del file {file_path}: {e}")
        return None

def save_file_xlsx(data, file_path):
    # verifica che l'estensione del file sia .xlsx
    if not file_path.endswith('.xlsx'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    if data is None:
        print(f"Errore: data è None, non posso salvare il file.")
        return None
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active

        for row in data:
            sheet.append(row)

        workbook.save(file_path)
        print(f"File salvato come {file_path}")
    except Exception as e:
        print(f"Errore durante il salvataggio del file {file_path}: {e}")

def load_file_mp3(file_path):
    # verifica che l'estensione del file sia .mp3
    if not file_path.endswith('.mp3'):
        print(f"Estensione del file {file_path} non supportata.")
        return None
    try:
        audio = AudioSegment.from_mp3(file_path)
        return audio
    except Exception as e:
        print(f"Errore durante la lettura del file {file_path}: {e}")
        return None